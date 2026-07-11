# admin_backend.py
# ==============================================================
# EMPOSIM 2026 — Python FastAPI Admin Backend
# Prerequisites: pip install fastapi uvicorn mysql-connector-python pandas pdfkit python-docx
# Run with: uvicorn admin_backend:app --reload --port 8000
# ==============================================================

from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import mysql.connector
import pandas as pd
import pdfkit
import docx
import os

app = FastAPI()

# Enable CORS for frontend integration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Database Configuration (Must match php/db.php)
DB_CONFIG = {
    'host': 'localhost',
    'database': 'emposim',
    'user': 'root',
    'password': ''
}

def get_db_connection():
    try:
        return mysql.connector.connect(**DB_CONFIG)
    except mysql.connector.Error as err:
        print(f"Error: {err}")
        return None

class UpdateRegistrationModel(BaseModel):
    name: str
    email: str
    phone: str
    college: str
    status: str

@app.get("/api/registrations")
def get_registrations():
    conn = get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")
    
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT id, student_name as name, email, phone, college, event_topic as event, payment_screenshot as screenshot, status FROM registrations ORDER BY id DESC")
    rows = cursor.fetchall()
    conn.close()
    return rows

@app.put("/api/registrations/{reg_id}")
def update_registration(reg_id: int, payload: UpdateRegistrationModel):
    conn = get_db_connection()
    if not conn:
        raise HTTPException(status_code=500, detail="Database connection failed")
    
    cursor = conn.cursor()
    query = ""\"
        UPDATE registrations 
        SET student_name = %s, email = %s, phone = %s, college = %s, status = %s 
        WHERE id = %s
    ""\"
    values = (payload.name, payload.email, payload.phone, payload.college, payload.status, reg_id)
    
    try:
        cursor.execute(query, values)
        conn.commit()
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()
        
    return {"success": True, "message": "Record updated successfully"}

# ----------------- EXPORTS -----------------

@app.get("/api/export/pdf")
def export_pdf():
    conn = get_db_connection()
    df = pd.read_sql("SELECT id, student_name, email, college, event_topic, status FROM registrations", con=conn)
    conn.close()
    
    # Generate HTML from Pandas
    html_content = ""\"
    <html>
    <head>
        <style>
            table { width: 100%; border-collapse: collapse; font-family: sans-serif; }
            th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            th { background-color: #f2f2f2; }
            h1 { font-family: sans-serif; }
        </style>
    </head>
    <body>
        <h1>Emposim 2026 - Registration Report</h1>
        {}
    </body>
    </html>
    ""\".format(df.to_html(index=False))
    
    # Convert HTML to PDF (Requires wkhtmltopdf installed on system)
    output_path = "export_registrations.pdf"
    try:
        pdfkit.from_string(html_content, output_path)
        return FileResponse(output_path, media_type='application/pdf', filename='Emposim_Registrations.pdf')
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed. Is wkhtmltopdf installed? Error: {str(e)}")

@app.get("/api/export/docx")
def export_docx():
    conn = get_db_connection()
    df = pd.read_sql("SELECT id, student_name, email, college, event_topic, status FROM registrations", con=conn)
    conn.close()
    
    doc = docx.Document()
    doc.add_heading('Emposim 2026 - Registration Report', 0)
    
    # Add Table
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = 'Table Grid'
    
    # Add Headers
    hdr_cells = t.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        
    # Add Rows
    for _, row in df.iterrows():
        row_cells = t.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            
    output_path = "export_registrations.docx"
    doc.save(output_path)
    
    return FileResponse(output_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename='Emposim_Registrations.docx')

